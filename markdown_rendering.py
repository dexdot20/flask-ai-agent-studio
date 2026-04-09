from __future__ import annotations

import re
from html import escape as html_escape

from docx import Document
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, Spacer, Table, TableStyle

_MARKDOWN_FENCE_RE = re.compile(r"^\s*```([A-Za-z0-9_-]+)?\s*$")
_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_MARKDOWN_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_MARKDOWN_ORDERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_MARKDOWN_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
_MARKDOWN_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_MARKDOWN_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
_MARKDOWN_HORIZONTAL_RULE_RE = re.compile(r"^\s*(?:[-*_]\s*){3,}$")


def _normalize_line_endings(text: str) -> str:
    return str(text or "").replace("\r\n", "\n").replace("\r", "\n")


def _escape_pdf_text(value: str) -> str:
    return html_escape(str(value or ""), quote=False)


def _find_math_delimiter(text: str, start_index: int, delimiter: str) -> int:
    search_index = start_index
    while search_index < len(text):
        delimiter_index = text.find(delimiter, search_index)
        if delimiter_index < 0:
            return -1
        if delimiter_index > start_index and text[delimiter_index - 1] == "\\":
            search_index = delimiter_index + len(delimiter)
            continue
        if delimiter == "$" and "\n" in text[start_index:delimiter_index]:
            search_index = delimiter_index + len(delimiter)
            continue
        return delimiter_index
    return -1


def _split_inline_math_segments(text: str) -> list[dict[str, object]]:
    value = _normalize_line_endings(text)
    if not value:
        return []
    if "$" not in value:
        return [{"type": "text", "text": value}]

    segments: list[dict[str, object]] = []
    buffer: list[str] = []
    index = 0

    def flush_buffer() -> None:
        if buffer:
            segments.append({"type": "text", "text": "".join(buffer)})
            buffer.clear()

    while index < len(value):
        char = value[index]
        if char == "\\":
            next_char = value[index + 1] if index + 1 < len(value) else ""
            if next_char == "$":
                buffer.append("$")
                index += 2
                continue
            buffer.append(char)
            index += 1
            continue

        if char != "$":
            buffer.append(char)
            index += 1
            continue

        display_mode = index + 1 < len(value) and value[index + 1] == "$"
        delimiter = "$$" if display_mode else "$"
        math_start = index + len(delimiter)
        math_end = _find_math_delimiter(value, math_start, delimiter)
        if math_end < 0:
            buffer.append(char)
            index += 1
            continue

        math_text = value[math_start:math_end].strip()
        if not math_text:
            buffer.append(delimiter)
            index = math_start
            continue

        flush_buffer()
        segments.append({"type": "math", "text": math_text, "display": display_mode})
        index = math_end + len(delimiter)

    flush_buffer()
    return segments


_PDF_TABLE_TOTAL_WIDTH = 450.0  # pts, approx A4 content area with standard margins
_PDF_INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__", re.DOTALL)
_PDF_INLINE_ITALIC_RE = re.compile(
    r"(?<!\w)\*(?!\s)(.+?)(?<!\s)\*(?!\w)|(?<!\w)_(?!\s)(.+?)(?<!\s)_(?!\w)", re.DOTALL
)
_PDF_INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _apply_pdf_inline_formatting(text: str) -> str:
    """Convert inline markdown formatting to ReportLab paragraph XML markup."""
    # Escape HTML entities first so bold/italic regexes only match raw markers
    value = html_escape(str(text or ""), quote=False)
    value = _PDF_INLINE_BOLD_RE.sub(lambda m: f"<b>{m.group(1) or m.group(2)}</b>", value)
    value = _PDF_INLINE_ITALIC_RE.sub(lambda m: f"<i>{m.group(1) or m.group(2)}</i>", value)
    value = _PDF_INLINE_CODE_RE.sub(lambda m: f'<font face="Courier">{m.group(1)}</font>', value)
    return value


def _render_pdf_inline_markup(text: str) -> str:
    parts: list[str] = []
    for segment in _split_inline_math_segments(text):
        if str(segment.get("type") or "") == "math":
            math_text = _escape_pdf_text(str(segment.get("text") or ""))
            parts.append(f'<font face="Courier">{math_text}</font>')
        else:
            parts.append(_apply_pdf_inline_formatting(str(segment.get("text") or "")))
    return "".join(parts)


def _append_docx_inline_runs(paragraph, text: str) -> None:
    segments = _split_inline_math_segments(text)
    if not segments:
        paragraph.add_run(" ")
        return

    for segment in segments:
        run = paragraph.add_run(str(segment.get("text") or ""))
        if str(segment.get("type") or "") == "math":
            run.italic = True
            run.font.name = "Courier New"


def _clean_markdown_inline(text: str, *, preserve_formatting: bool = False) -> str:
    value = _normalize_line_endings(text)

    def _replace_image(match: re.Match) -> str:
        alt_text = (match.group(1) or "").strip()
        url = (match.group(2) or "").strip()
        return alt_text or url

    def _replace_link(match: re.Match) -> str:
        label = _clean_markdown_inline(match.group(1) or "", preserve_formatting=preserve_formatting)
        url = (match.group(2) or "").strip()
        if not label:
            return url
        if label == url:
            return label
        return f"{label} ({url})" if url else label

    value = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", _replace_image, value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", _replace_link, value)
    if not preserve_formatting:
        value = re.sub(r"`([^`]+)`", lambda match: match.group(1), value)
        value = re.sub(r"(\*\*|__)(.+?)\1", lambda match: match.group(2), value)
        value = re.sub(r"(?<!\w)(\*|_)(?!\s)(.+?)(?<!\s)\1(?!\w)", lambda match: match.group(2), value)
    value = re.sub(r"~~(.+?)~~", lambda match: match.group(1), value)
    value = re.sub(r"<(https?://[^>]+)>", lambda match: match.group(1), value)
    value = value.replace("\\*", "*").replace("\\_", "_").replace("\\`", "`")
    value = value.replace("\\[", "[").replace("\\]", "]").replace("\\(", "(").replace("\\)", ")")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _iter_markdown_blocks(text: str, *, preserve_inline_formatting: bool = False) -> list[dict[str, object]]:
    normalized = _normalize_line_endings(text)
    if not normalized.strip():
        return []

    blocks: list[dict[str, object]] = []
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    list_kind: str | None = None
    table_rows: list[list[str]] = []
    code_lines: list[str] = []
    in_code = False
    code_language = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        paragraph_text = _clean_markdown_inline(
            " ".join(part.strip() for part in paragraph_lines if part.strip()),
            preserve_formatting=preserve_inline_formatting,
        )
        paragraph_lines = []
        if paragraph_text:
            blocks.append({"type": "paragraph", "text": paragraph_text})

    def flush_list() -> None:
        nonlocal list_items, list_kind
        if not list_items:
            return
        blocks.append({"type": "list", "kind": list_kind or "bullet", "items": list_items[:]})
        list_items = []
        list_kind = None

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        if preserve_inline_formatting:
            cleaned: list[list[str]] = [
                [_clean_markdown_inline(cell or "", preserve_formatting=True) for cell in row]
                for row in table_rows
            ]
            blocks.append({"type": "table", "rows": cleaned})
        else:
            for row in table_rows:
                row_text = _clean_markdown_inline(" | ".join(cell for cell in row if cell is not None))
                if row_text:
                    blocks.append({"type": "paragraph", "text": row_text})
        table_rows = []

    for line in normalized.split("\n"):
        fence_match = _MARKDOWN_FENCE_RE.match(line)
        if fence_match:
            if in_code:
                block_text = "\n".join(code_lines)
                if code_language in {"markdown", "md", "mkd"}:
                    flush_paragraph()
                    flush_list()
                    flush_table()
                    blocks.extend(_iter_markdown_blocks(block_text, preserve_inline_formatting=preserve_inline_formatting))
                else:
                    blocks.append({"type": "code", "text": block_text})
                code_lines = []
                code_language = ""
                in_code = False
            else:
                flush_paragraph()
                flush_list()
                flush_table()
                in_code = True
                code_language = (fence_match.group(1) or "").strip().lower()
            continue

        if in_code:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            flush_list()
            flush_table()
            continue

        if _MARKDOWN_HORIZONTAL_RULE_RE.match(line):
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append({"type": "spacer"})
            continue

        heading_match = _MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": _clean_markdown_inline(heading_match.group(2), preserve_formatting=preserve_inline_formatting),
                }
            )
            continue

        if _MARKDOWN_TABLE_SEPARATOR_RE.match(line):
            continue

        if _MARKDOWN_TABLE_ROW_RE.match(line) and line.strip().count("|") >= 2:
            flush_paragraph()
            flush_list()
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            continue

        bullet_match = _MARKDOWN_BULLET_RE.match(line)
        if bullet_match:
            flush_paragraph()
            flush_table()
            if list_kind not in {None, "bullet"}:
                flush_list()
            list_kind = "bullet"
            list_items.append(_clean_markdown_inline(bullet_match.group(1), preserve_formatting=preserve_inline_formatting))
            continue

        ordered_match = _MARKDOWN_ORDERED_RE.match(line)
        if ordered_match:
            flush_paragraph()
            flush_table()
            if list_kind not in {None, "ordered"}:
                flush_list()
            list_kind = "ordered"
            list_items.append(_clean_markdown_inline(ordered_match.group(2), preserve_formatting=preserve_inline_formatting))
            continue

        quote_match = _MARKDOWN_QUOTE_RE.match(line)
        if quote_match:
            flush_paragraph()
            flush_list()
            flush_table()
            quote_text = _clean_markdown_inline(quote_match.group(1), preserve_formatting=preserve_inline_formatting)
            if quote_text:
                blocks.append({"type": "paragraph", "text": quote_text})
            continue

        flush_table()
        paragraph_lines.append(line)

    if in_code:
        block_text = "\n".join(code_lines)
        if code_language in {"markdown", "md", "mkd"}:
            blocks.extend(_iter_markdown_blocks(block_text, preserve_inline_formatting=preserve_inline_formatting))
        else:
            blocks.append({"type": "code", "text": block_text})

    flush_paragraph()
    flush_list()
    flush_table()
    return blocks


def append_markdown_pdf_story(
    story: list,
    markdown_text: str,
    *,
    body_style,
    heading1_style,
    heading_style,
    subheading_style,
    code_style,
    heading_level_offset: int = 0,
    empty_text: str = "(empty)",
) -> None:
    blocks = _iter_markdown_blocks(markdown_text)
    if not blocks:
        story.append(Paragraph(_escape_pdf_text(empty_text), body_style))
        return

    offset = max(0, int(heading_level_offset or 0))
    for block in blocks:
        block_type = str(block.get("type") or "")
        if block_type == "heading":
            level = min(6, max(1, int(block.get("level") or 1) + offset))
            style = heading1_style if level <= 1 else heading_style if level == 2 else subheading_style
            story.append(Paragraph(_render_pdf_inline_markup(str(block.get("text") or "Untitled")), style))
        elif block_type == "paragraph":
            paragraph_text = str(block.get("text") or "").strip()
            if paragraph_text:
                story.append(Paragraph(_render_pdf_inline_markup(paragraph_text), body_style))
        elif block_type == "list":
            items = [str(item).strip() for item in block.get("items") or [] if str(item).strip()]
            if items:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(_render_pdf_inline_markup(item), body_style)) for item in items],
                        bulletType="1" if str(block.get("kind") or "bullet") == "ordered" else "bullet",
                        leftIndent=18,
                    )
                )
                story.append(Spacer(1, 4))
        elif block_type == "code":
            code_text = str(block.get("text") or "").rstrip()
            if code_text:
                story.append(Preformatted(code_text, code_style))
                story.append(Spacer(1, 4))
        elif block_type == "table":
            table_rows_data = block.get("rows") or []
            if table_rows_data:
                num_cols = max((len(r) for r in table_rows_data), default=1)
                col_width = _PDF_TABLE_TOTAL_WIDTH / num_cols
                table_cell_style = ParagraphStyle(
                    "TableCell",
                    parent=body_style,
                    fontSize=max(7, (getattr(body_style, "fontSize", 10) or 10) - 1),
                    leading=max(9, (getattr(body_style, "leading", 14) or 14) - 2),
                    spaceAfter=2,
                    spaceBefore=2,
                )
                table_data = [
                    [
                        Paragraph(_render_pdf_inline_markup(str(cell or "")), table_cell_style)
                        for cell in row
                    ]
                    for row in table_rows_data
                ]
                rendered_table = Table(
                    table_data, colWidths=[col_width] * num_cols, hAlign="LEFT", repeatRows=0
                )
                rendered_table.setStyle(
                    TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c8cdd8")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2fa")),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fd")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ])
                )
                story.append(rendered_table)
                story.append(Spacer(1, 8))
        elif block_type == "spacer":
            story.append(Spacer(1, 4))


def append_markdown_docx(
    document: Document,
    markdown_text: str,
    *,
    heading_level_offset: int = 0,
    empty_text: str = "(empty)",
) -> None:
    blocks = _iter_markdown_blocks(markdown_text)
    if not blocks:
        document.add_paragraph(empty_text)
        return

    offset = max(0, int(heading_level_offset or 0))
    for block in blocks:
        block_type = str(block.get("type") or "")
        if block_type == "heading":
            level = min(9, max(1, int(block.get("level") or 1) + offset))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _append_docx_inline_runs(paragraph, str(block.get("text") or "Untitled"))
        elif block_type == "paragraph":
            paragraph_text = str(block.get("text") or "").strip()
            paragraph = document.add_paragraph()
            _append_docx_inline_runs(paragraph, paragraph_text or " ")
        elif block_type == "list":
            style_name = "List Number" if str(block.get("kind") or "bullet") == "ordered" else "List Bullet"
            for item in block.get("items") or []:
                item_text = str(item).strip()
                if not item_text:
                    continue
                try:
                    paragraph = document.add_paragraph(style=style_name)
                    _append_docx_inline_runs(paragraph, item_text)
                except Exception:
                    prefix = "1. " if style_name == "List Number" else "- "
                    paragraph = document.add_paragraph()
                    _append_docx_inline_runs(paragraph, f"{prefix}{item_text}")
        elif block_type == "code":
            code_text = str(block.get("text") or "").rstrip()
            if code_text:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(code_text)
                run.font.name = "Courier New"
        elif block_type == "spacer":
            document.add_paragraph(" ")