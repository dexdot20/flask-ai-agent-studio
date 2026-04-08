from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from markdown_rendering import append_markdown_docx, append_markdown_pdf_story

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from canvas_service import extract_canvas_documents
from db import extract_message_attachments, extract_sub_agent_traces

_FONT_PATHS = {
    "DejaVuSans": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "DejaVuSans-Bold": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "DejaVuSansMono": "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
}


def _try_register_fonts() -> bool:
    import os

    if not all(os.path.exists(path) for path in _FONT_PATHS.values()):
        return False
    try:
        for name, path in _FONT_PATHS.items():
            pdfmetrics.registerFont(TTFont(name, path))
        return True
    except Exception:
        return False


_UNICODE_FONTS = _try_register_fonts()
_BODY_FONT = "DejaVuSans" if _UNICODE_FONTS else "Helvetica"
_BOLD_FONT = "DejaVuSans-Bold" if _UNICODE_FONTS else "Helvetica-Bold"
_MONO_FONT = "DejaVuSansMono" if _UNICODE_FONTS else "Courier"


def _escape_pdf_text(value: str) -> str:
    return str(value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _safe_text(value, fallback: str = "") -> str:
    text = str(value or "").strip()
    return text or fallback


def _append_detail_line(lines: list[str], label: str, value) -> None:
    text = _safe_text(value)
    if text:
        lines.append(f"- {label}: {text}")


def _append_boolean_detail_line(lines: list[str], label: str, value) -> None:
    if value is True:
        lines.append(f"- {label}: Yes")
    elif value is False:
        lines.append(f"- {label}: No")


def _join_values(values: list[str]) -> str:
    return ", ".join(value for value in values if value)


def _build_export_header(conversation: dict, message_count: int) -> list[str]:
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    model = str(conversation.get("model") or "").strip()
    conversation_id = conversation.get("id")
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    lines = [f"# {title}", "", f"Exported at: {exported_at}"]
    if conversation_id is not None:
        lines.append(f"Conversation ID: {conversation_id}")
    lines.append(f"Message count: {message_count}")
    if model:
        lines.append(f"Model: {model}")
    return lines


def _build_message_metadata_details(message: dict) -> str | None:
    lines = []
    _append_detail_line(lines, "Message ID", message.get("id"))
    _append_detail_line(lines, "Position", message.get("position"))
    _append_detail_line(lines, "Created at", message.get("created_at"))
    _append_detail_line(lines, "Tool call ID", message.get("tool_call_id"))
    _append_detail_line(lines, "Prompt tokens", message.get("prompt_tokens"))
    _append_detail_line(lines, "Completion tokens", message.get("completion_tokens"))
    _append_detail_line(lines, "Total tokens", message.get("total_tokens"))

    usage = message.get("usage") if isinstance(message.get("usage"), dict) else None
    if usage:
        usage_lines = []
        _append_detail_line(usage_lines, "Prompt tokens", usage.get("prompt_tokens"))
        _append_detail_line(usage_lines, "Completion tokens", usage.get("completion_tokens"))
        _append_detail_line(usage_lines, "Total tokens", usage.get("total_tokens"))
        _append_detail_line(usage_lines, "Estimated input tokens", usage.get("estimated_input_tokens"))
        if usage_lines:
            lines.extend(["", "Usage"])
            lines.extend(usage_lines)

    return "\n".join(lines).strip() or None


def _build_canvas_documents_details(canvas_documents: list[dict]) -> str | None:
    lines = []
    for document in canvas_documents:
        if not isinstance(document, dict):
            continue
        title = _safe_text(document.get("title"), fallback="Canvas")
        lines.append(f"### Canvas: {title}")
        doc_content = _safe_text(document.get("content"))
        if doc_content:
            lines.extend(["", "```markdown", doc_content, "```", ""])
    return "\n".join(lines).strip() or None


def _build_attachment_details(attachments: list[dict]) -> str | None:
    blocks = []
    for index, attachment in enumerate(attachments, start=1):
        if not isinstance(attachment, dict):
            continue

        kind = _safe_text(attachment.get("kind"), fallback="attachment")
        title = _safe_text(
            attachment.get("file_name")
            or attachment.get("image_name")
            or attachment.get("video_title")
            or attachment.get("file_id")
            or attachment.get("image_id")
            or attachment.get("video_id"),
            fallback="Attachment",
        )
        heading = f"#### Attachment {index}: {kind.title()}"
        if title and title.lower() != kind.lower():
            heading = f"{heading} - {title}"

        lines = [heading]
        if kind == "image":
            _append_detail_line(lines, "Image ID", attachment.get("image_id"))
            _append_detail_line(lines, "Name", attachment.get("image_name"))
            _append_detail_line(lines, "MIME type", attachment.get("image_mime_type"))
            _append_detail_line(lines, "Analysis method", attachment.get("analysis_method"))
            _append_detail_line(lines, "OCR text", attachment.get("ocr_text"))
            _append_detail_line(lines, "Vision summary", attachment.get("vision_summary"))
            _append_detail_line(lines, "Assistant guidance", attachment.get("assistant_guidance"))
            key_points = attachment.get("key_points") if isinstance(attachment.get("key_points"), list) else []
            if key_points:
                lines.append("- Key points:")
                for point in key_points[:8]:
                    point_text = _safe_text(point)
                    if point_text:
                        lines.append(f"  - {point_text}")
        elif kind == "video":
            _append_detail_line(lines, "Video ID", attachment.get("video_id"))
            _append_detail_line(lines, "Title", attachment.get("video_title"))
            _append_detail_line(lines, "URL", attachment.get("video_url"))
            _append_detail_line(lines, "Platform", attachment.get("video_platform"))
            _append_detail_line(lines, "Transcript language", attachment.get("transcript_language"))
            _append_detail_line(lines, "Transcript context", attachment.get("transcript_context_block"))
            _append_boolean_detail_line(lines, "Transcript truncated", attachment.get("transcript_text_truncated"))
        else:
            _append_detail_line(lines, "File ID", attachment.get("file_id"))
            _append_detail_line(lines, "File name", attachment.get("file_name"))
            _append_detail_line(lines, "MIME type", attachment.get("file_mime_type"))
            _append_detail_line(lines, "Submission mode", attachment.get("submission_mode"))
            _append_detail_line(lines, "Canvas mode", attachment.get("canvas_mode"))
            _append_detail_line(lines, "File context", attachment.get("file_context_block"))
            _append_boolean_detail_line(lines, "File truncated", attachment.get("file_text_truncated"))
            page_ids = attachment.get("visual_page_image_ids") if isinstance(attachment.get("visual_page_image_ids"), list) else []
            if page_ids:
                _append_detail_line(lines, "Visual page image IDs", _join_values([_safe_text(page_id) for page_id in page_ids[:8]]))
            _append_detail_line(lines, "Visual page count", attachment.get("visual_page_count"))

        blocks.append("\n".join(lines).strip())
    return "\n\n".join(blocks).strip() or None


def _build_sub_agent_trace_details(traces: list[dict]) -> str | None:
    blocks = []
    for index, trace in enumerate(traces, start=1):
        if not isinstance(trace, dict):
            continue

        task = _safe_text(trace.get("task"), fallback=_safe_text(trace.get("task_full"), fallback="Sub-agent trace"))
        lines = [f"#### Trace {index}: {task}"]
        _append_detail_line(lines, "Task full", trace.get("task_full"))
        _append_detail_line(lines, "Status", trace.get("status"))
        _append_detail_line(lines, "Model", trace.get("model"))
        _append_detail_line(lines, "Summary", trace.get("summary"))
        _append_detail_line(lines, "Reasoning", trace.get("reasoning"))
        _append_detail_line(lines, "Fallback note", trace.get("fallback_note"))
        _append_detail_line(lines, "Error", trace.get("error"))
        _append_boolean_detail_line(lines, "Canvas saved", trace.get("canvas_saved"))
        canvas_document_id = _safe_text(trace.get("canvas_document_id"))
        canvas_document_title = _safe_text(trace.get("canvas_document_title"))
        if canvas_document_id or canvas_document_title:
            if canvas_document_id and canvas_document_title:
                lines.append(f"- Canvas document: {canvas_document_title} ({canvas_document_id})")
            else:
                lines.append(f"- Canvas document: {_safe_text(canvas_document_title or canvas_document_id)}")

        tool_trace = trace.get("tool_trace") if isinstance(trace.get("tool_trace"), list) else []
        if tool_trace:
            lines.append("")
            lines.append("##### Tool Trace")
            for step_index, entry in enumerate(tool_trace[:8], start=1):
                if not isinstance(entry, dict):
                    continue
                step = entry.get("step")
                step_label = f"Step {int(step)}" if isinstance(step, (int, float)) else f"Step {step_index}"
                tool_name = _safe_text(entry.get("tool_name"), fallback="tool")
                summary = _safe_text(entry.get("summary") or entry.get("preview"))
                line = f"- {step_label}: {tool_name}"
                if summary:
                    line = f"{line}: {summary}"
                lines.append(line)

        artifacts = trace.get("artifacts") if isinstance(trace.get("artifacts"), list) else []
        if artifacts:
            lines.append("")
            lines.append("##### Artifacts")
            for artifact_index, artifact in enumerate(artifacts[:8], start=1):
                if not isinstance(artifact, dict):
                    continue
                artifact_kind = _safe_text(artifact.get("kind"), fallback="artifact").title()
                artifact_label = _safe_text(artifact.get("label") or artifact.get("name") or f"Artifact {artifact_index}")
                artifact_value = _safe_text(artifact.get("value") or artifact.get("content") or artifact.get("summary"))
                line = f"- {artifact_kind}: {artifact_label}"
                if artifact_value:
                    line = f"{line} — {artifact_value}"
                lines.append(line)

        messages = trace.get("messages") if isinstance(trace.get("messages"), list) else []
        if messages:
            lines.append("")
            lines.append("##### Messages")
            for message_index, message in enumerate(messages[:8], start=1):
                if not isinstance(message, dict):
                    continue
                role = _safe_text(message.get("role"), fallback="message").title()
                content = _safe_text(message.get("content") or message.get("text"))
                if content:
                    lines.append(f"- {message_index}. {role}: {content}")
                else:
                    lines.append(f"- {message_index}. {role}")

        blocks.append("\n".join(lines).strip())
    return "\n\n".join(blocks).strip() or None


def _iter_message_sections(messages: list[dict]) -> list[dict]:
    sections = []
    section_index = 0
    for message in messages or []:
        if not isinstance(message, dict):
            continue

        role = str(message.get("role") or "message").strip() or "message"
        content = str(message.get("content") or "").strip()
        metadata = message.get("metadata") if isinstance(message.get("metadata"), dict) else {}
        details = []

        message_metadata = _build_message_metadata_details(message)
        if message_metadata:
            details.append(("Message Metadata", message_metadata))

        tool_trace = metadata.get("tool_trace") if isinstance(metadata.get("tool_trace"), list) else []
        if tool_trace:
            lines = []
            for entry in tool_trace:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                step = entry.get("step")
                summary = str(entry.get("summary") or entry.get("preview") or "").strip()
                prefix = f"Step {int(step)}" if isinstance(step, (int, float)) else "Step"
                lines.append(f"- {prefix}: {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Trace", "\n".join(lines)))

        tool_results = metadata.get("tool_results") if isinstance(metadata.get("tool_results"), list) else []
        if tool_results:
            lines = []
            for entry in tool_results:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                summary = str(entry.get("summary") or "").strip()
                lines.append(f"- {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Results", "\n".join(lines)))

        attachments = extract_message_attachments(metadata)
        if attachments:
            attachment_details = _build_attachment_details(attachments)
            if attachment_details:
                details.append(("Attachments", attachment_details))

        sub_agent_traces = extract_sub_agent_traces(metadata)
        if sub_agent_traces:
            trace_details = _build_sub_agent_trace_details(sub_agent_traces)
            if trace_details:
                details.append(("Sub-Agent Traces", trace_details))

        canvas_documents = extract_canvas_documents(metadata)
        if canvas_documents:
            canvas_details = _build_canvas_documents_details(canvas_documents)
            if canvas_details:
                details.append(("Canvas Documents", canvas_details))

        has_user_facing_details = any(label != "Message Metadata" for label, _ in details)
        if role == "assistant" and not content and not has_user_facing_details:
            continue

        section_index += 1

        sections.append(
            {
                "title": f"## {section_index}. {role.title()}",
                "content": content,
                "details": details,
            }
        )
    return sections


def build_conversation_markdown_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    lines = _build_export_header(conversation, len(sections))
    for section in sections:
        lines.extend(["", section["title"], ""])
        lines.append(section["content"] or "_(empty)_")
        for label, value in section["details"]:
            lines.extend(["", f"### {label}", "", value])
    lines.append("")
    return "\n".join(lines).encode("utf-8")


def build_conversation_docx_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    document = Document()
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    document.add_heading(title, level=0)
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    meta_parts = [f"Exported at: {exported_at}", f"Message count: {len(sections)}"]
    conversation_id = conversation.get("id")
    if conversation_id is not None:
        meta_parts.insert(1, f"Conversation ID: {conversation_id}")
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    document.add_paragraph(" | ".join(meta_parts))

    for section in sections:
        document.add_heading(section["title"].replace("## ", ""), level=1)
        append_markdown_docx(document, section["content"] or "(empty)", heading_level_offset=2)
        for label, value in section["details"]:
            document.add_heading(label, level=2)
            append_markdown_docx(document, str(value or "").strip() or "(empty)", heading_level_offset=2)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_conversation_pdf_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ConversationTitle", parent=styles["Title"], fontName=_BOLD_FONT)
    heading_style = ParagraphStyle(
        "ConversationHeading",
        parent=styles["Heading2"],
        fontName=_BOLD_FONT,
        textColor=colors.HexColor("#1f2a44"),
        spaceAfter=8,
        spaceBefore=12,
    )
    subheading_style = ParagraphStyle(
        "ConversationSubheading",
        parent=styles["Heading3"],
        fontName=_BOLD_FONT,
        textColor=colors.HexColor("#33415f"),
        spaceAfter=6,
        spaceBefore=8,
    )
    body_style = ParagraphStyle(
        "ConversationBody",
        parent=styles["BodyText"],
        fontName=_BODY_FONT,
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "ConversationCode",
        parent=styles["Code"],
        fontName=_MONO_FONT,
        fontSize=8.5,
        leading=11,
        leftIndent=10,
        rightIndent=10,
        backColor=colors.HexColor("#f3f5f9"),
        borderPadding=8,
    )

    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    story = [Paragraph(_escape_pdf_text(title), title_style), Spacer(1, 6)]
    meta_parts = [f"Exported at: {exported_at}", f"Message count: {len(sections)}"]
    conversation_id = conversation.get("id")
    if conversation_id is not None:
        meta_parts.insert(1, f"Conversation ID: {conversation_id}")
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    story.append(Paragraph(_escape_pdf_text(" | ".join(meta_parts)), body_style))

    for section in sections:
        story.append(Spacer(1, 6))
        story.append(Paragraph(_escape_pdf_text(section["title"].replace("## ", "")), heading_style))
        append_markdown_pdf_story(
            story,
            section["content"] or "(empty)",
            body_style=body_style,
            heading1_style=subheading_style,
            heading_style=subheading_style,
            subheading_style=subheading_style,
            code_style=code_style,
            heading_level_offset=2,
        )
        for label, value in section["details"]:
            story.append(Paragraph(_escape_pdf_text(label), subheading_style))
            append_markdown_pdf_story(
                story,
                str(value or "").strip() or "(empty)",
                body_style=body_style,
                heading1_style=subheading_style,
                heading_style=subheading_style,
                subheading_style=subheading_style,
                code_style=code_style,
                heading_level_offset=2,
            )

    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    doc.build(story)
    return output.getvalue()